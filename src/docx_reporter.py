#!/usr/bin/env python3
"""
docx_reporter.py - Gerador de Relatórios DOCX EXPANDIDO
Sistema de Análise Estatística DL 54/2018 v2.0

VERSÃO EXPANDIDA: ~60 páginas com 15+ secções detalhadas

Autor: Sistema de Qualidade de Dados Educacionais
"""

import logging
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from config import ConfigManager
from utils import format_percentage, format_number, create_timestamp


logger = logging.getLogger('DL54.docx_reporter')


class DocxReporter:
    """
    Gerador de relatórios DOCX profissionais EXPANDIDO.
    
    Características:
    - Página de título
    - Índice
    - 15+ secções detalhadas
    - Tabelas formatadas
    - Gráficos incorporados
    - Estilo consistente
    - ~60 páginas de análises
    """
    
    def __init__(self, config: ConfigManager, logger_instance: logging.Logger):
        """Inicializa gerador de relatórios DOCX expandido."""
        self.config = config
        self.logger = logger_instance
        
        self.report_config = config.get_report_config()
        self.personalization = config.get_personalization()
        self.escola_mapping = config.get_escola_mapping()
        
        self.output_dir = Path(config.get('IO', 'OUTPUT_DIR'))
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def generate_report(self, stats: Dict[str, Any],
                       chart_paths: Optional[List[Path]] = None) -> Path:
        """Gera relatório DOCX completo EXPANDIDO."""
        self.logger.info("=" * 80)
        self.logger.info("GERANDO RELATÓRIO DOCX EXPANDIDO")
        self.logger.info("=" * 80)
        
        doc = Document()
        self._setup_styles(doc)
        
        # 1. Página de título
        self._add_title_page(doc)
        doc.add_page_break()
        
        # 2. Enquadramento legal
        if self.personalization.get('include_legal_framework', True):
            self._add_legal_framework(doc)
            doc.add_page_break()
        
        # 3. Análise Global
        if 'global' in stats:
            self._add_global_analysis(doc, stats['global'])
            doc.add_page_break()
        
        # 4. Análise por Escola
        if 'por_escola' in stats:
            self._add_school_analysis(doc, stats['por_escola'])
            doc.add_page_break()
        
        # 5. Análise por Ano
        if 'por_ano' in stats:
            self._add_year_analysis(doc, stats['por_ano'])
            doc.add_page_break()
        
        # NOVA: 6. Análise por Turma
        if 'por_turma' in stats:
            self._add_turma_analysis(doc, stats['por_turma'])
            doc.add_page_break()
        
        # NOVA: 7. Análise Ano + Turma
        if 'por_ano_turma' in stats:
            self._add_ano_turma_analysis(doc, stats['por_ano_turma'])
            doc.add_page_break()
        
        # NOVA: 8. Estatísticas por Aluno (Turma)
        if 'estatisticas_aluno_turma' in stats:
            self._add_estatisticas_aluno_analysis(doc, stats['estatisticas_aluno_turma'])
            doc.add_page_break()
        
        # NOVA: 9. Alíneas Detalhadas
        if 'alineas_detalhadas' in stats:
            self._add_alineas_detalhadas_analysis(doc, stats['alineas_detalhadas'])
            doc.add_page_break()
        
        # NOVA: 10. Terapias Completas
        if 'terapias_completas' in stats:
            self._add_terapias_analysis(doc, stats['terapias_completas'])
            doc.add_page_break()
        
        # 11. Análise por Sexo (base)
        if 'por_sexo' in stats:
            self._add_gender_analysis(doc, stats['por_sexo'])
            doc.add_page_break()
        
        # NOVA: 12. Sexo Detalhado (todas alíneas)
        if 'sexo_detalhado' in stats:
            self._add_sexo_detalhado_analysis(doc, stats['sexo_detalhado'])
            doc.add_page_break()
        
        # 13. Análise por Escalão ASE
        if 'por_escalao_ase' in stats:
            self._add_ase_analysis(doc, stats['por_escalao_ase'])
            doc.add_page_break()
        
        # 14. Rankings
        if 'rankings' in stats:
            self._add_rankings(doc, stats['rankings'])
            doc.add_page_break()
        
        # 15. Gráficos
        if chart_paths and self.report_config['charts']:
            self._add_charts_section(doc, chart_paths)
        
        # Salvar
        timestamp = create_timestamp()
        filename = self.report_config['docx_filename'].replace('.docx', f'_EXPANDIDO_{timestamp}.docx')
        output_path = self.output_dir / filename
        doc.save(output_path)
        
        self.logger.info(f"✓ Relatório DOCX EXPANDIDO gerado: {output_path.name}")
        self.logger.info("=" * 80)
        
        return output_path
    
    def _setup_styles(self, doc: Document):
        """Configura estilos do documento."""
        styles = doc.styles
        
        try:
            h1_style = styles['Heading 1']
            h1_style.font.size = Pt(16)
            h1_style.font.color.rgb = RGBColor(68, 114, 196)
            h1_style.font.bold = True
        except:
            pass
        
        try:
            h2_style = styles['Heading 2']
            h2_style.font.size = Pt(14)
            h2_style.font.color.rgb = RGBColor(0, 153, 204)
            h2_style.font.bold = True
        except:
            pass
    
    def _create_table(self, doc: Document, headers: List[str]):
        """Helper: cria tabela com headers."""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Light Grid Accent 1'
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        return table
    
    def _add_table_row(self, table, values: List[str]):
        """Helper: adiciona linha à tabela."""
        row_cells = table.add_row().cells
        for i, value in enumerate(values):
            row_cells[i].text = str(value)
    
    def _add_title_page(self, doc: Document):
        """Adiciona página de título."""
        self.logger.info("Criando página de título...")
        
        for _ in range(5):
            doc.add_paragraph()
        
        title = doc.add_paragraph()
        title_run = title.add_run(self.personalization.get('title', 'Análise Estatística EXPANDIDA'))
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(68, 114, 196)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if self.personalization.get('subtitle'):
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(self.personalization['subtitle'])
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.italic = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for _ in range(3):
            doc.add_paragraph()
        
        if self.personalization.get('organization'):
            org = doc.add_paragraph()
            org_run = org.add_run(self.personalization['organization'])
            org_run.font.size = Pt(14)
            org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        date = doc.add_paragraph()
        date_run = date.add_run(datetime.now().strftime('%d de %B de %Y'))
        date_run.font.size = Pt(12)
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if self.personalization.get('version'):
            version = doc.add_paragraph()
            version_run = version.add_run(f"Versão {self.personalization['version']} - EXPANDIDA")
            version_run.font.size = Pt(10)
            version_run.font.italic = True
            version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _add_legal_framework(self, doc: Document):
        """Adiciona enquadramento legal."""
        self.logger.info("Adicionando enquadramento legal...")
        
        doc.add_heading('1. Enquadramento Legal', level=1)
        
        doc.add_paragraph(
            'Este relatório apresenta uma análise estatística EXPANDIDA das medidas de apoio '
            'à aprendizagem e à inclusão, no âmbito do Decreto-Lei n.º 54/2018, '
            'de 6 de julho, que estabelece o regime jurídico da educação inclusiva.'
        )
        
        doc.add_heading('Medidas de Apoio à Aprendizagem e à Inclusão', level=2)
        
        doc.add_paragraph(
            'As medidas de apoio à aprendizagem e à inclusão agrupam-se em três níveis '
            'de intervenção: universais, seletivas e adicionais.'
        )
        
        doc.add_paragraph('Medidas Universais (Artigo 8.º):', style='List Bullet')
        doc.add_paragraph(
            'Correspondem às respostas educativas que a escola tem disponíveis para todos os alunos.',
            style='List Bullet 2'
        )
        
        doc.add_paragraph('Medidas Seletivas (Artigo 9.º):', style='List Bullet')
        doc.add_paragraph(
            'Visam colmatar as necessidades de suporte à aprendizagem não supridas pela aplicação '
            'de medidas universais.',
            style='List Bullet 2'
        )
        
        doc.add_paragraph('Medidas Adicionais (Artigo 10.º):', style='List Bullet')
        doc.add_paragraph(
            'Visam colmatar dificuldades acentuadas e persistentes ao nível da comunicação, '
            'interação, cognição ou aprendizagem.',
            style='List Bullet 2'
        )
    
    def _add_global_analysis(self, doc: Document, global_stats: Dict):
        """Adiciona análise global."""
        self.logger.info("Adicionando análise global...")
        
        doc.add_heading('2. Análise Global', level=1)
        
        total_alunos = global_stats['total_alunos']
        
        doc.add_paragraph(
            f'A análise incide sobre um universo de {format_number(total_alunos)} alunos.'
        )
        
        # Medidas principais
        doc.add_heading('2.1. Medidas Principais', level=2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Medida'
        hdr_cells[1].text = 'N'
        hdr_cells[2].text = '%'
        hdr_cells[3].text = 'Sem Medida'
        
        for medida, dados in global_stats['medidas_principais'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = medida
            row_cells[1].text = format_number(dados['n'])
            row_cells[2].text = format_percentage(dados['percentagem'])
            row_cells[3].text = format_number(dados.get('sem_medida', 0))
        
        doc.add_paragraph()
        
        # Medidas detalhadas
        doc.add_heading('2.2. Medidas Detalhadas', level=2)
        
        for tipo_nome, tipo_key in [
            ('Universais (Art. 8º)', 'universais'),
            ('Seletivas (Art. 9º)', 'seletivas'),
            ('Adicionais (Art. 10º)', 'adicionais')
        ]:
            doc.add_heading(f'2.2.{tipo_key[0].upper()}. {tipo_nome}', level=3)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Medida'
            hdr_cells[1].text = 'N'
            hdr_cells[2].text = '%'
            
            for medida, dados in global_stats['medidas_detalhadas'][tipo_key].items():
                row_cells = table.add_row().cells
                row_cells[0].text = medida
                row_cells[1].text = format_number(dados['n'])
                row_cells[2].text = format_percentage(dados['percentagem'])
            
            doc.add_paragraph()
    
    def _add_school_analysis(self, doc: Document, school_stats: Dict):
        """Adiciona análise por escola."""
        self.logger.info("Adicionando análise por escola...")
        
        doc.add_heading('3. Análise por Escola', level=1)
        
        doc.add_paragraph(
            f'Análise detalhada de {len(school_stats)} escolas do agrupamento.'
        )
        
        for codigo, escola_dados in school_stats.items():
            doc.add_heading(f"3.{codigo}. {escola_dados['nome']}", level=2)
            
            doc.add_paragraph(
                f"Total de alunos: {format_number(escola_dados['total_alunos'])} "
                f"({format_percentage(escola_dados['peso_no_agrupamento'])} do agrupamento)"
            )
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Medida'
            hdr_cells[1].text = 'N'
            hdr_cells[2].text = '%'
            
            for medida, dados in escola_dados['medidas_principais'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = medida
                row_cells[1].text = format_number(dados['n'])
                row_cells[2].text = format_percentage(dados['percentagem'])
            
            doc.add_paragraph()
    
    def _add_year_analysis(self, doc: Document, year_stats: Dict):
        """Adiciona análise por ano."""
        self.logger.info("Adicionando análise por ano...")
        
        doc.add_heading('4. Análise por Ano de Escolaridade', level=1)
        
        doc.add_paragraph(
            f'Distribuição das medidas pelos {len(year_stats)} anos de escolaridade.'
        )
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ano'
        hdr_cells[1].text = 'Total Alunos'
        hdr_cells[2].text = 'M. Universais'
        hdr_cells[3].text = 'M. Seletivas'
        hdr_cells[4].text = 'M. Adicionais'
        
        for ano in sorted(year_stats.keys()):
            ano_dados = year_stats[ano]
            row_cells = table.add_row().cells
            
            row_cells[0].text = ano_dados['ano_nome']
            row_cells[1].text = format_number(ano_dados['total_alunos'])
            
            medidas = ano_dados['medidas_principais']
            row_cells[2].text = f"{medidas.get('Medidas Universais', {}).get('n', 0)}"
            row_cells[3].text = f"{medidas.get('Medidas Seletivas', {}).get('n', 0)}"
            row_cells[4].text = f"{medidas.get('Medidas Adicionais', {}).get('n', 0)}"
    
    def _add_turma_analysis(self, doc: Document, turma_stats: Dict):
        """NOVA: Adiciona análise por turma."""
        self.logger.info("Adicionando análise por turma...")
        
        doc.add_heading('5. Análise por Turma', level=1)
        
        doc.add_paragraph(
            f'Análise detalhada das {len(turma_stats)} turmas existentes no agrupamento.'
        )
        
        # Tabela resumo
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Turma'
        hdr_cells[1].text = 'Total Alunos'
        hdr_cells[2].text = 'M. Universais'
        hdr_cells[3].text = 'M. Seletivas'
        hdr_cells[4].text = 'M. Adicionais'
        
        for turma, dados in sorted(turma_stats.items()):
            row_cells = table.add_row().cells
            row_cells[0].text = turma
            row_cells[1].text = format_number(dados['total_alunos'])
            
            medidas = dados['medidas_principais']
            row_cells[2].text = f"{medidas.get('Medidas Universais', {}).get('n', 0)}"
            row_cells[3].text = f"{medidas.get('Medidas Seletivas', {}).get('n', 0)}"
            row_cells[4].text = f"{medidas.get('Medidas Adicionais', {}).get('n', 0)}"
    
    def _add_ano_turma_analysis(self, doc: Document, ano_turma_stats: Dict):
        """NOVA: Adiciona análise ano + turma."""
        self.logger.info("Adicionando análise ano + turma...")
        
        doc.add_heading('6. Análise por Ano e Turma', level=1)
        
        doc.add_paragraph(
            f'Análise detalhada de {len(ano_turma_stats)} combinações ano + turma.'
        )
        
        # Agrupar por ano
        anos_dict = {}
        for chave, dados in ano_turma_stats.items():
            ano = dados['ano']
            if ano not in anos_dict:
                anos_dict[ano] = []
            anos_dict[ano].append((chave, dados))
        
        for ano in sorted(anos_dict.keys()):
            ano_nome = "Pré-Escolar" if ano == 0 else f"{ano}º Ano"
            doc.add_heading(f'6.{ano}. {ano_nome}', level=2)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Turma'
            hdr_cells[1].text = 'Total'
            hdr_cells[2].text = 'M. Univ.'
            hdr_cells[3].text = 'M. Selet.'
            hdr_cells[4].text = 'M. Adic.'
            
            for chave, dados in sorted(anos_dict[ano], key=lambda x: x[1]['turma']):
                row_cells = table.add_row().cells
                row_cells[0].text = dados['turma']
                row_cells[1].text = format_number(dados['total_alunos'])
                
                medidas = dados['medidas_principais']
                row_cells[2].text = f"{medidas.get('Medidas Universais', {}).get('n', 0)}"
                row_cells[3].text = f"{medidas.get('Medidas Seletivas', {}).get('n', 0)}"
                row_cells[4].text = f"{medidas.get('Medidas Adicionais', {}).get('n', 0)}"
            
            doc.add_paragraph()
    
    def _add_estatisticas_aluno_analysis(self, doc: Document, stats_aluno: Dict):
        """NOVA: Adiciona estatísticas por aluno dentro de cada turma."""
        self.logger.info("Adicionando estatísticas por aluno...")
        
        doc.add_heading('7. Estatísticas por Aluno (dentro de cada Turma)', level=1)
        
        doc.add_paragraph(
            'Análise do número de medidas POR ALUNO dentro de cada turma: '
            'máximo, mínimo, média, mediana, desvio padrão e percentis.'
        )
        
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Turma'
        hdr_cells[1].text = 'N Alunos'
        hdr_cells[2].text = 'Máx'
        hdr_cells[3].text = 'Mín'
        hdr_cells[4].text = 'Média'
        hdr_cells[5].text = 'Mediana'
        hdr_cells[6].text = 'Desv.Pad'
        hdr_cells[7].text = 'P25-P75'
        
        for chave, dados in sorted(stats_aluno.items()):
            row_cells = table.add_row().cells
            stats_medidas = dados['estatisticas_medidas_por_aluno']
            
            row_cells[0].text = dados['nome']
            row_cells[1].text = format_number(dados['total_alunos'])
            row_cells[2].text = str(stats_medidas['maximo'])
            row_cells[3].text = str(stats_medidas['minimo'])
            row_cells[4].text = f"{stats_medidas['media']:.2f}"
            row_cells[5].text = f"{stats_medidas['mediana']:.1f}"
            row_cells[6].text = f"{stats_medidas['desvio_padrao']:.2f}"
            row_cells[7].text = f"{stats_medidas['p25']:.1f}-{stats_medidas['p75']:.1f}"
    
    def _add_alineas_detalhadas_analysis(self, doc: Document, alineas_stats: Dict):
        """NOVA: Adiciona análise detalhada de alíneas."""
        self.logger.info("Adicionando análise de alíneas...")
        
        doc.add_heading('8. Análise Detalhada por Alíneas', level=1)
        
        doc.add_paragraph(
            'Análise de todas as 15 alíneas das medidas (5 universais, 5 seletivas, 5 adicionais) '
            'discriminadas por ano de escolaridade.'
        )
        
        for ano_nome, dados_ano in sorted(alineas_stats.items()):
            doc.add_heading(f'8.X. {ano_nome}', level=2)
            doc.add_paragraph(f"Total de alunos: {format_number(dados_ano['total_alunos'])}")
            
            for tipo_nome, tipo_key in [
                ('Medidas Universais', 'universais'),
                ('Medidas Seletivas', 'seletivas'),
                ('Medidas Adicionais', 'adicionais')
            ]:
                if dados_ano[tipo_key]:
                    doc.add_heading(f'{tipo_nome}', level=3)
                    
                    table = doc.add_table(rows=1, cols=3)
                    table.style = 'Light Grid Accent 1'
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Alínea'
                    hdr_cells[1].text = 'N'
                    hdr_cells[2].text = '%'
                    
                    for alinea, dados in dados_ano[tipo_key].items():
                        row_cells = table.add_row().cells
                        row_cells[0].text = alinea
                        row_cells[1].text = format_number(dados['n'])
                        row_cells[2].text = format_percentage(dados['percentagem'])
                    
                    doc.add_paragraph()
    
    def _add_terapias_analysis(self, doc: Document, terapias_stats: Dict):
        """NOVA: Adiciona análise de terapias."""
        self.logger.info("Adicionando análise de terapias...")
        
        doc.add_heading('9. Análise de Terapias', level=1)
        
        doc.add_paragraph(
            'Análise das 5 terapias principais: Terapia da Fala, Terapia Ocupacional, '
            'Psicomotricidade, Fisioterapia e Psicologia.'
        )
        
        # Global
        doc.add_heading('9.1. Terapias - Visão Global', level=2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Terapia'
        hdr_cells[1].text = 'N'
        hdr_cells[2].text = '%'
        
        for terapia, dados in terapias_stats['global'].items():
            row_cells = table.add_row().cells
            row_cells[0].text = terapia
            row_cells[1].text = format_number(dados['n'])
            row_cells[2].text = format_percentage(dados['percentagem'])
        
        doc.add_paragraph()
        
        # Por ano
        doc.add_heading('9.2. Terapias por Ano de Escolaridade', level=2)
        
        for ano_nome, terapias_ano in sorted(terapias_stats['por_ano'].items()):
            doc.add_heading(f'{ano_nome}', level=3)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Terapia'
            hdr_cells[1].text = 'N'
            hdr_cells[2].text = '%'
            
            for terapia, dados in terapias_ano.items():
                row_cells = table.add_row().cells
                row_cells[0].text = terapia
                row_cells[1].text = format_number(dados['n'])
                row_cells[2].text = format_percentage(dados['percentagem'])
            
            doc.add_paragraph()
        
        # Por sexo
        doc.add_heading('9.3. Terapias por Sexo', level=2)
        
        for sexo, terapias_sexo in sorted(terapias_stats['por_sexo'].items()):
            doc.add_heading(f'Sexo {sexo}', level=3)
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Terapia'
            hdr_cells[1].text = 'N'
            hdr_cells[2].text = '%'
            
            for terapia, dados in terapias_sexo.items():
                row_cells = table.add_row().cells
                row_cells[0].text = terapia
                row_cells[1].text = format_number(dados['n'])
                row_cells[2].text = format_percentage(dados['percentagem'])
    
    def _add_gender_analysis(self, doc: Document, gender_stats: Dict):
        """Adiciona análise por sexo (base)."""
        self.logger.info("Adicionando análise por sexo...")
        
        doc.add_heading('10. Análise por Sexo', level=1)
        
        sexos = {k: v for k, v in gender_stats.items() if not k.startswith('_')}
        
        for sexo, dados in sexos.items():
            doc.add_heading(f'10.{sexo}. Sexo {sexo}', level=2)
            
            doc.add_paragraph(
                f"Total: {format_number(dados['total_alunos'])} alunos "
                f"({format_percentage(dados['percentagem_do_total'])} do total)"
            )
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Medida'
            hdr_cells[1].text = 'N'
            hdr_cells[2].text = '% no Grupo'
            
            for medida, medida_dados in dados['medidas_principais'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = medida
                row_cells[1].text = format_number(medida_dados['n'])
                row_cells[2].text = format_percentage(medida_dados['percentagem_no_grupo'])
            
            doc.add_paragraph()
    
    def _add_sexo_detalhado_analysis(self, doc: Document, sexo_stats: Dict):
        """NOVA: Adiciona análise detalhada por sexo (todas alíneas)."""
        self.logger.info("Adicionando análise detalhada por sexo...")
        
        doc.add_heading('11. Análise Detalhada por Sexo (Todas as Alíneas)', level=1)
        doc.add_paragraph('Análise discriminada de todas as 15 alíneas por sexo (M/F).')
        
        for sexo, dados in sorted(sexo_stats.items()):
            doc.add_heading(f'11.{sexo}. Sexo {sexo}', level=2)
            doc.add_paragraph(f"Total de alunos: {format_number(dados['total_alunos'])}")
            
            for tipo_nome, tipo_key in [('Medidas Universais', 'universais'),
                                         ('Medidas Seletivas', 'seletivas'),
                                         ('Medidas Adicionais', 'adicionais')]:
                doc.add_heading(f'{tipo_nome}', level=3)
                table = self._create_table(doc, ['Alínea', 'N', '%'])
                for alinea, alinea_dados in dados['medidas_detalhadas'][tipo_key].items():
                    self._add_table_row(table, [alinea, format_number(alinea_dados['n']),
                                                format_percentage(alinea_dados['percentagem'])])
                doc.add_paragraph()
    
    def _add_ase_analysis(self, doc: Document, ase_stats: Dict):
        """Adiciona análise por escalão ASE."""
        self.logger.info("Adicionando análise por escalão ASE...")
        
        doc.add_heading('12. Análise por Escalão ASE', level=1)
        
        doc.add_paragraph(
            'Distribuição das medidas por escalão socioeconómico (Ação Social Escolar).'
        )
        
        for escalao, dados in ase_stats.items():
            doc.add_heading(f'12.{escalao}. Escalão {escalao}', level=2)
            
            doc.add_paragraph(
                f"Total: {format_number(dados['total_alunos'])} alunos "
                f"({format_percentage(dados['percentagem_do_total'])} do total)"
            )
            
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Medida'
            hdr_cells[1].text = 'N'
            hdr_cells[2].text = '%'
            
            for medida, medida_dados in dados['medidas_principais'].items():
                row_cells = table.add_row().cells
                row_cells[0].text = medida
                row_cells[1].text = format_number(medida_dados['n'])
                row_cells[2].text = format_percentage(medida_dados['percentagem'])
            
            doc.add_paragraph()
    
    def _add_rankings(self, doc: Document, rankings: Dict):
        """Adiciona rankings."""
        self.logger.info("Adicionando rankings...")
        
        doc.add_heading('13. Rankings por Medida', level=1)
        
        doc.add_paragraph(
            'Classificação das escolas por percentagem de alunos com cada tipo de medida.'
        )
        
        for medida, ranking_data in rankings.items():
            doc.add_heading(f'13.{medida}. {medida}', level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Posição'
            hdr_cells[1].text = 'Escola'
            hdr_cells[2].text = 'N'
            hdr_cells[3].text = '%'
            
            for item in ranking_data:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item['rank'])
                row_cells[1].text = item['nome'].split(',')[0]
                row_cells[2].text = format_number(item['alunos_com_medida'])
                row_cells[3].text = format_percentage(item['percentagem'])
            
            doc.add_paragraph()
    
    def _add_charts_section(self, doc: Document, chart_paths: List[Path]):
        """Adiciona secção de gráficos."""
        self.logger.info("Adicionando gráficos ao relatório...")
        
        doc.add_heading('14. Representações Gráficas', level=1)
        
        for i, chart_path in enumerate(chart_paths, 1):
            if chart_path.exists():
                doc.add_heading(f'14.{i}. {chart_path.stem.replace("_", " ").title()}', level=2)
                
                try:
                    doc.add_picture(str(chart_path), width=Inches(6))
                    doc.add_paragraph()
                except Exception as e:
                    self.logger.warning(f"Não foi possível adicionar gráfico: {e}")
